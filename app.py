from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from datetime import datetime
import os
import pymorphy3

# Получаем сегодняшнюю дату
сегодня = datetime.today().date()

# Форматируем дату в нужном формате (например, ДД.ММ.ГГГГ)
formatted_today = сегодня.strftime("%d.%m.%Y")

app = Flask(__name__)
    
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Получаем данные из формы
        sudyamne = request.form['sudyamne']
        suduchastok = request.form['suduchastok']
        region = request.form['region']
        city = request.form['city']
        FIOMINI = request.form['FIOMINI']
        formatted_FIO = format_fio(FIOMINI)
        sudebprikaz = request.form['sudebprikaz']
        datasudebprikaz = request.form['datasudebprikaz']
        money = request.form['money']
        vodnik = request.form['vodnik']
        phone = request.form['phone']
        mail = request.form['mail']
        adress = request.form['adress']
        choiceSEX = request.form['choiceSEX']
        # Определяем род и некоторые фразы в зависимости от пола
        if choiceSEX == '1':
            uznal = 'узнал'
            biluved = 'был уведомлен'
            deliv = 'получал'
            imel = 'имел'
        elif choiceSEX == '2':
            uznal = 'узнала'
            biluved = 'была уведомлена'
            deliv = 'получала'
            imel = 'имела'
        else:
            return "Неправильный ввод пола"
        # Создание нового документа Word
        doc = Document()
        # Функция для добавления параграфа с заданными настройками
        def add_paragraph(text, alignment, bold=False, first_line_indent=0):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if bold:
                run.bold = True
            paragraph.alignment = alignment
            paragraph.paragraph_format.right_indent = Inches(-0.60)
            paragraph.paragraph_format.left_indent = Inches(-0.60)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
            paragraph.paragraph_format.line_spacing = 1.0
        # Добавляем информацию в документ
        add_paragraph(
            f"Мировому судье {suduchastok} судебного участка\n"
            f"{region}, г. {city}\n"
            f"{sudyamne}\n"
            f"Взыскатель: {vodnik},\n"
            f"Должник: {FIOMINI}\n" 
            f"адрес: {adress},\n"
            f"телефон: {phone},\n"
            f"адрес электронной почты: {mail},\n"
            f"Судебный приказ {sudebprikaz} от {datasudebprikaz}\n",
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        )

                # Добавление заголовка
        add_paragraph(
            f"В О З Р А Ж Е Н И Я\n"
            f"относительно исполнения судебного приказа с ходатайством о восстановлении\n"
            f"пропущенного процессуального срока",
            WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True
        )

        # Добавление параграфа с красной строкой
        add_paragraph(
            f"{datasudebprikaz} года мировым судьей судебного участка №{suduchastok} судебного района г.Кургана Курганской области вынесен судебный приказ о взыскании с должника ({formatted_FIO}) в пользу {vodnik} задолженности в сумме {money}₽. На основании указанного судебного приказа был наложен арест на банковские счета.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Судебный приказ считаю незаконным и необоснованным, указанные в нем суммы являются чрезмерно завышенными, возражаю относительно его исполнения.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Возражения заявлены мною за пределами срока, предусмотренного для подачи возражений относительно исполнения судебного приказа, полагаю, что следующие обстоятельства указывают на наличие уважительных причин, по которым мною пропущен срок подачи возражений относительно исполнения судебного приказа.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"О вынесенном в отношении меня судебном приказе я {uznal} {formatted_today}г. по факту ареста средств на картах. Я как  Должник своевременно не {biluved} о судебном производстве, возбужденном в отношении меня, а также о том, что Взыскатель взыскивает с меня какую бы то ни было задолженность. Копию судебного приказа {sudebprikaz} от {datasudebprikaz} я не {deliv}. Таким образом, я не {imel} возможности своевременно подать возражения относительно исполнения судебного приказа.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"На основании изложенного, руководствуясь ст.ст. 112, 128 ГПК РФ,",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"П Р О Ш У :",
            WD_PARAGRAPH_ALIGNMENT.LEFT,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Восстановить срок подачи возражений относительно исполнения судебного приказа {sudebprikaz} от {datasudebprikaz} мировым судьей судебного участка №{suduchastok} судебного района г. Кургана Курганской области о взыскании с {FIOMINI} в пользу {vodnik} задолженности в сумме {money}₽.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Отменить судебный приказ {sudebprikaz} от {datasudebprikaz} мировым судьей судебного участка №{suduchastok} судебного района г. Кургана Курганской области о взыскании с должника ({formatted_FIO}) в пользу {vodnik} задолженности в сумме {money}₽.",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Принятое по результатам рассмотрения данных возражений судебное постановление выручить мне лично после предварительного уведомления на номер сотового телефона: {phone}.\n",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )

        add_paragraph(
            f"Дата: {formatted_today}г.                                                      Подпись: ____________________",
            WD_PARAGRAPH_ALIGNMENT.LEFT,
            first_line_indent=0.45  # Установка отступа первой строки в 17 мм
        )
        # Сохранение документа
        file_path = 'document.docx'
        doc.save(file_path)
        return send_file(file_path, as_attachment=True)
    return render_template('index.html')



def format_fio(fio):
    parts = fio.split()
    if len(parts) < 3:
        return fio  # Возвращаем оригинал, если не хватает частей
    surname = parts[0]
    initials = '. '.join(part[0] for part in parts[1:]) + '.'
    return f"{surname} {initials}"
if __name__ == '__main__':
    app.run(debug=True)